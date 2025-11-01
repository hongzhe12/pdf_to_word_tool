import os
import sys

from PySide6.QtCore import QThread, Signal, Qt
from PySide6.QtGui import QFont, QTextCursor, QIcon
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout,
                               QHBoxLayout, QPushButton, QLabel, QTextEdit,
                               QFileDialog, QProgressBar, QMessageBox, QGroupBox,
                               QSpinBox, QFormLayout)
from docx.shared import Cm
import fitz
from docx import Document
from PIL import Image
import resources_rc

class ConversionThread(QThread):
    """转换线程"""
    progress_updated = Signal(int)
    message_updated = Signal(str)
    conversion_finished = Signal(bool)

    def __init__(self, pdf_folder, target_width, target_height, images_per_row):
        super().__init__()
        self.pdf_folder = pdf_folder
        self.target_width = target_width
        self.target_height = target_height
        self.images_per_row = images_per_row
        self.is_running = True

    def run(self):
        try:
            # 获取所有PDF文件
            pdf_files = [f for f in os.listdir(self.pdf_folder) if f.lower().endswith('.pdf')]
            total_files = len(pdf_files)

            if total_files == 0:
                self.message_updated.emit("在指定文件夹中未找到PDF文件！")
                self.conversion_finished.emit(False)
                return

            image_data_list = []
            for i, file in enumerate(pdf_files):
                if not self.is_running:
                    break

                input_file = os.path.join(self.pdf_folder, file)
                self.message_updated.emit(f"正在处理: {file}")

                image_data = self.convert_pdf_to_image(input_file)
                if image_data:
                    image_data_list.append((file, image_data))

                # 更新进度
                progress = int((i + 1) / total_files * 50)  # 转换阶段占50%
                self.progress_updated.emit(progress)

            if image_data_list and self.is_running:
                self.message_updated.emit("正在生成Word文档...")
                success = self.add_images_to_word(image_data_list)
                if success:
                    self.message_updated.emit(f"转换完成！共处理 {len(image_data_list)} 个文件")
                    self.conversion_finished.emit(True)
                else:
                    self.conversion_finished.emit(False)
            else:
                self.conversion_finished.emit(False)

        except Exception as e:
            self.message_updated.emit(f"转换过程中发生错误: {str(e)}")
            self.conversion_finished.emit(False)

    def convert_pdf_to_image(self, pdf_path, dpi=200):
        """将PDF转换为图片数据（不保存文件）"""
        try:
            # 打开PDF并转换为图片
            pdf_document = fitz.open(pdf_path)
            page = pdf_document[0]

            # 设置转换参数
            zoom = dpi / 72
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)

            # 转换为PIL Image并调整尺寸
            pil_image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # 毫米转像素并调整尺寸
            mm_to_pixel = lambda mm: int(mm * dpi / 25.4)
            target_px = (mm_to_pixel(self.target_width), mm_to_pixel(self.target_height))

            resized_image = pil_image.resize(target_px, Image.Resampling.LANCZOS)

            # 将图片保存到内存中
            from io import BytesIO
            img_buffer = BytesIO()
            resized_image.save(img_buffer, format='PNG', optimize=True)
            img_data = img_buffer.getvalue()
            img_buffer.close()

            pdf_document.close()
            self.message_updated.emit(f"处理成功: {os.path.basename(pdf_path)}")
            return img_data

        except Exception as e:
            self.message_updated.emit(f"处理失败 {os.path.basename(pdf_path)}: {e}")
            return None

    def add_images_to_word(self, image_data_list):
        """将图片数据直接添加到Word文档中"""
        try:
            word_path = os.path.join(self.pdf_folder, "火车票集合.docx")
            doc = Document()

            # 设置页面边距
            section = doc.sections[0]
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

            # 创建表格来排列图片
            table = doc.add_table(rows=0, cols=self.images_per_row)
            table.autofit = False

            row_cells = None
            for i, (filename, img_data) in enumerate(image_data_list):
                if i % self.images_per_row == 0:
                    row_cells = table.add_row().cells

                cell = row_cells[i % self.images_per_row]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = 1  # 居中对齐

                # 添加图片到单元格，设置宽度为6cm
                run = paragraph.add_run()

                # 从内存数据添加图片
                from io import BytesIO
                img_buffer = BytesIO(img_data)
                run.add_picture(img_buffer, width=Cm(6))
                img_buffer.close()

                # 更新进度（Word生成阶段占50%）
                progress = 50 + int((i + 1) / len(image_data_list) * 50)
                self.progress_updated.emit(progress)

            doc.save(word_path)
            self.message_updated.emit(f"Word文档已保存至: {word_path}")
            return True

        except Exception as e:
            self.message_updated.emit(f"生成Word文档失败: {e}")
            return False

    def stop(self):
        self.is_running = False


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.pdf_folder = ""
        self.init_ui()
        self.conversion_thread = None

    def init_ui(self):
        self.setWindowTitle("PDF转Word工具 - 火车票处理")
        self.setFixedSize(600, 650)  # 调整窗口高度

        # 中央部件
        central_widget = QWidget()
        self.setCentralWidget(central_widget)

        # 主布局
        layout = QVBoxLayout(central_widget)

        # 标题
        title_label = QLabel("PDF转Word工具")
        title_label.setAlignment(Qt.AlignCenter)
        title_font = QFont()
        title_font.setPointSize(16)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)

        # 参数设置组
        params_group = QGroupBox("参数设置")
        params_layout = QFormLayout()

        self.width_spin = QSpinBox()
        self.width_spin.setRange(10, 200)
        self.width_spin.setValue(75)
        self.width_spin.setSuffix(" mm")

        self.height_spin = QSpinBox()
        self.height_spin.setRange(10, 200)
        self.height_spin.setValue(45)
        self.height_spin.setSuffix(" mm")

        self.images_per_row_spin = QSpinBox()
        self.images_per_row_spin.setRange(1, 5)
        self.images_per_row_spin.setValue(2)
        self.images_per_row_spin.setSuffix(" 张/行")

        params_layout.addRow("图片宽度:", self.width_spin)
        params_layout.addRow("图片高度:", self.height_spin)
        params_layout.addRow("每行图片数:", self.images_per_row_spin)

        params_group.setLayout(params_layout)
        layout.addWidget(params_group)

        # 文件夹选择组
        folder_group = QGroupBox("文件夹设置")
        folder_layout = QVBoxLayout()

        # PDF文件夹选择
        pdf_layout = QHBoxLayout()
        self.pdf_folder_label = QLabel("未选择PDF文件夹")
        self.pdf_folder_label.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.pdf_folder_btn = QPushButton("选择PDF文件夹")
        self.pdf_folder_btn.clicked.connect(self.select_pdf_folder)
        pdf_layout.addWidget(self.pdf_folder_label, 1)
        pdf_layout.addWidget(self.pdf_folder_btn)
        folder_layout.addLayout(pdf_layout)

        folder_group.setLayout(folder_layout)
        layout.addWidget(folder_group)

        # 进度条
        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        layout.addWidget(self.progress_bar)

        # 控制按钮
        button_layout = QHBoxLayout()
        self.start_btn = QPushButton("开始转换")
        self.start_btn.clicked.connect(self.start_conversion)
        self.start_btn.setEnabled(False)

        self.stop_btn = QPushButton("停止转换")
        self.stop_btn.clicked.connect(self.stop_conversion)
        self.stop_btn.setEnabled(False)

        button_layout.addWidget(self.start_btn)
        button_layout.addWidget(self.stop_btn)
        layout.addLayout(button_layout)

        # 日志输出
        log_group = QGroupBox("转换日志")
        log_layout = QVBoxLayout()
        self.log_text = QTextEdit()
        self.log_text.setReadOnly(True)
        log_layout.addWidget(self.log_text)
        log_group.setLayout(log_layout)
        layout.addWidget(log_group)

        # 状态栏
        self.status_label = QLabel("就绪")
        layout.addWidget(self.status_label)

    def select_pdf_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "选择PDF文件夹")
        if folder:
            self.pdf_folder = folder
            self.pdf_folder_label.setText(folder)
            self.start_btn.setEnabled(True)

    def start_conversion(self):
        if not self.pdf_folder:
            QMessageBox.warning(self, "警告", "请先选择PDF文件夹！")
            return

        # 禁用开始按钮，启用停止按钮
        self.start_btn.setEnabled(False)
        self.stop_btn.setEnabled(True)
        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)

        # 清空日志
        self.log_text.clear()

        # 创建转换线程
        self.conversion_thread = ConversionThread(
            self.pdf_folder,
            self.width_spin.value(),
            self.height_spin.value(),
            self.images_per_row_spin.value()
        )

        # 连接信号
        self.conversion_thread.progress_updated.connect(self.progress_bar.setValue)
        self.conversion_thread.message_updated.connect(self.add_log_message)
        self.conversion_thread.conversion_finished.connect(self.conversion_finished)

        # 启动线程
        self.conversion_thread.start()
        self.status_label.setText("正在转换...")

    def stop_conversion(self):
        if self.conversion_thread and self.conversion_thread.isRunning():
            self.conversion_thread.stop()
            self.conversion_thread.wait()
            self.add_log_message("转换已停止")
            self.status_label.setText("转换已停止")

        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)

    def conversion_finished(self, success):
        self.start_btn.setEnabled(True)
        self.stop_btn.setEnabled(False)
        self.status_label.setText("转换完成" if success else "转换失败")

        if success:
            QMessageBox.information(self, "完成", "转换完成！Word文档已生成在PDF文件夹中。")
        else:
            QMessageBox.warning(self, "完成", "转换失败！")

    def add_log_message(self, message):
        self.log_text.append(f"{message}")
        # 自动滚动到底部
        cursor = self.log_text.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        self.log_text.setTextCursor(cursor)

    def closeEvent(self, event):
        if self.conversion_thread and self.conversion_thread.isRunning():
            self.conversion_thread.stop()
            self.conversion_thread.wait()
        event.accept()


if __name__ == "__main__":
    from ctypes import windll  # Only exists on Windows.
    from PySide6.QtGui import QIcon
    myappid = 'mycompany.myproduct.subproduct.version'
    windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon("images/icon.svg"))
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
